from __future__ import annotations

from pathlib import Path

import pytest
from docx import Document
from pypdf import PdfWriter

from spec_viewer.errors import SpecViewerError
from spec_viewer.normalization import normalize_document, sha256_file
from spec_viewer.schema import load_document
from spec_viewer.workspace import init_workspace


def _workspace(repository: Path, name: str = "Normalization Demo") -> Path:
    relative = init_workspace(name, repository / "workspaces", repository)
    return repository / relative


def test_normalizes_markdown_without_changing_source(
    isolated_repository: Path, tmp_path: Path
) -> None:
    source = tmp_path / "spec.md"
    source.write_bytes(b"# Demo\r\n\r\nRequirement with spaces.   \r\n")
    source_hash = sha256_file(source)
    workspace = _workspace(isolated_repository)

    relative = normalize_document(workspace, source, isolated_repository)

    assert relative.as_posix() == "normalized/specification.md"
    assert (workspace / relative).read_text() == "# Demo\n\nRequirement with spaces.\n"
    assert sha256_file(source) == source_hash
    assert sha256_file(workspace / "source/spec.md") == source_hash
    project = load_document(workspace / "project.yaml")
    assert project["status"]["lifecycle"] == "normalized"
    assert project["source"]["original_file"] == "source/spec.md"


def test_normalizes_docx_headings_lists_and_table(
    isolated_repository: Path, tmp_path: Path
) -> None:
    source = tmp_path / "spec.docx"
    document = Document()
    document.add_heading("Demo specification", level=1)
    document.add_paragraph("First requirement")
    document.add_paragraph("List item", style="List Bullet")
    table = document.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "Field"
    table.cell(0, 1).text = "Value"
    table.cell(1, 0).text = "Role"
    table.cell(1, 1).text = "Owner"
    document.save(str(source))

    workspace = _workspace(isolated_repository, "DOCX Demo")
    normalize_document(workspace, source, isolated_repository)
    markdown = (workspace / "normalized/specification.md").read_text()

    assert "# Demo specification" in markdown
    assert "- List item" in markdown
    assert "| Field | Value |" in markdown


def test_rejects_scanned_or_empty_pdf(isolated_repository: Path, tmp_path: Path) -> None:
    source = tmp_path / "scan.pdf"
    writer = PdfWriter()
    writer.add_blank_page(width=100, height=100)
    with source.open("wb") as stream:
        writer.write(stream)
    with pytest.raises(SpecViewerError, match="OCR is unsupported"):
        normalize_document(_workspace(isolated_repository, "PDF Demo"), source, isolated_repository)


def test_rejects_duplicate_and_unsupported_source(
    isolated_repository: Path, tmp_path: Path
) -> None:
    source = tmp_path / "spec.md"
    source.write_text("# Demo", encoding="utf-8")
    workspace = _workspace(isolated_repository, "Duplicate Demo")
    normalize_document(workspace, source, isolated_repository)
    with pytest.raises(SpecViewerError, match="already exists"):
        normalize_document(workspace, source, isolated_repository)

    unsupported = tmp_path / "spec.txt"
    unsupported.write_text("Demo", encoding="utf-8")
    with pytest.raises(SpecViewerError, match="Unsupported"):
        normalize_document(
            _workspace(isolated_repository, "TXT Demo"), unsupported, isolated_repository
        )
